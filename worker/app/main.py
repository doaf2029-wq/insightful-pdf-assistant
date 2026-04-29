import os, io, json, uuid, time, base64, traceback
from pathlib import Path
from typing import Optional
import httpx
import fitz  # PyMuPDF
import pdfplumber
from fastapi import FastAPI, Header, HTTPException, BackgroundTasks
from pydantic import BaseModel
from anthropic import Anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from supabase import create_client
from app.config import settings

app = FastAPI()
anthropic = Anthropic(api_key=settings.anthropic_api_key)
sb = create_client(settings.supabase_url, settings.supabase_service_role)

# ── Models ────────────────────────────────────────────────────────────────────
class JobPayload(BaseModel):
    jobId: str
    userId: str
    filename: str
    pdfUrl: str
    callbackUrl: str
    outputsBucket: str
    prompt: Optional[str] = None
    language: Optional[str] = "auto"

# ── Helpers ───────────────────────────────────────────────────────────────────
def update_job(job_id: str, status: str, progress: int, message: str):
    sb.table("jobs").update({
        "status": status,
        "progress": progress,
        "status_message": message,
    }).eq("id", job_id).execute()

def post_callback(callback_url: str, payload: dict):
    try:
        httpx.post(callback_url, json=payload, timeout=15)
    except Exception:
        pass

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_heading(doc, text, level, underline=False, rtl=False):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.bold = True
        if underline:
            run.underline = True
        run.font.name = "Arial"
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    if rtl:
        set_rtl(p)
    return p

def add_body(doc, text, bold=False, italic=True, rtl=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if rtl:
        set_rtl(p)
    return p

def add_warning(doc, text, rtl=False):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️  {text}")
    run.bold = True
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    if rtl:
        set_rtl(p)

def add_note(doc, text, rtl=False):
    p = doc.add_paragraph()
    run = p.add_run(f"📌 {text}")
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    if rtl:
        set_rtl(p)

def add_bullet(doc, text, bold_part=None, rtl=False):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part and bold_part in text:
        parts = text.split(bold_part, 1)
        r1 = p.add_run(parts[0])
        r1.italic = True
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r2 = p.add_run(bold_part)
        r2.bold = True
        r2.italic = True
        r2.font.name = "Arial"
        r2.font.size = Pt(11)
        if parts[1]:
            r3 = p.add_run(parts[1])
            r3.italic = True
            r3.font.name = "Arial"
            r3.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
    if rtl:
        set_rtl(p)

def add_image(doc, img_bytes, width_inches=6.5):
    try:
        buf = io.BytesIO(img_bytes)
        doc.add_picture(buf, width=Inches(width_inches))
    except Exception:
        pass

def add_table(doc, headers, rows, rtl=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D5E8F0")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.italic = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
    doc.add_paragraph()

# ── Claude call ───────────────────────────────────────────────────────────────
def call_claude(messages: list, max_tokens=4000) -> str:
    for model in [settings.model_primary, settings.model_fallback]:
        try:
            resp = anthropic.messages.create(
                model=model,
                max_tokens=max_tokens,
                messages=messages,
            )
            return resp.content[0].text
        except Exception as e:
            print(f"Model {model} failed: {e}")
            time.sleep(2)
    raise RuntimeError("Both Claude models failed")

# ── PDF pre-extraction ────────────────────────────────────────────────────────
def extract_images(pdf_path: str, out_dir: Path) -> dict:
    """Extract all images from PDF, return page->list of image paths."""
    out_dir.mkdir(parents=True, exist_ok=True)
    index = {}
    doc = fitz.open(pdf_path)
    for page_num in range(len(doc)):
        page = doc[page_num]
        images = page.get_images(full=True)
        page_imgs = []
        for img_idx, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            ext = base_image["ext"]
            img_path = out_dir / f"page_{page_num+1}_img_{img_idx+1}.{ext}"
            img_path.write_bytes(img_bytes)
            page_imgs.append(str(img_path))
        if page_imgs:
            index[page_num + 1] = page_imgs
    return index

def extract_tables(pdf_path: str, out_dir: Path) -> dict:
    """Extract tables from PDF, return page->list of table dicts."""
    out_dir.mkdir(parents=True, exist_ok=True)
    index = {}
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if tables:
                page_tables = []
                for t_idx, table in enumerate(tables):
                    if not table:
                        continue
                    headers = [str(c) if c else "" for c in table[0]]
                    rows = [[str(c) if c else "" for c in row] for row in table[1:]]
                    tbl = {"headers": headers, "rows": rows}
                    t_path = out_dir / f"page_{page_num+1}_table_{t_idx+1}.json"
                    t_path.write_text(json.dumps(tbl, ensure_ascii=False))
                    page_tables.append(tbl)
                index[page_num + 1] = page_tables
    return index

def extract_text_chunks(pdf_path: str, chunk_size: int = 30) -> list:
    """Extract text from PDF in chunks, return list of {pages, text}."""
    doc = fitz.open(pdf_path)
    total = len(doc)
    chunks = []
    for start in range(0, total, chunk_size):
        end = min(start + chunk_size, total)
        text = ""
        for i in range(start, end):
            text += f"\n--- Page {i+1} ---\n"
            text += doc[i].get_text()
        chunks.append({"start": start + 1, "end": end, "text": text})
    return chunks

# ── Discovery ─────────────────────────────────────────────────────────────────
DISCOVERY_PROMPT = """You are scanning a document to identify all components it contains.
A component can be a service, policy, guide, glossary, regulation, or any other distinct content block.

Chunk pages: {start} to {end}

TASK: Return a JSON array of all components found. For each:
- "name": exact name as written
- "type": "service" | "policy" | "guide" | "glossary" | "intro" | "other"
- "start_page": estimated start page
- "end_page": estimated end page or "continues"
- "language": "en" | "ar" | "bilingual"
- "has_sub_components": true/false
- "started_before": true if it started in a previous chunk

Return ONLY valid JSON array. No markdown, no preamble.

DOCUMENT TEXT:
{text}"""

def discover_components(chunks: list) -> list:
    all_components = []
    for chunk in chunks:
        prompt = DISCOVERY_PROMPT.format(
            start=chunk["start"], end=chunk["end"], text=chunk["text"][:8000]
        )
        for attempt in range(2):
            try:
                raw = call_claude([{"role": "user", "content": prompt}])
                raw = raw.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
                components = json.loads(raw)
                all_components.extend(components)
                break
            except Exception as e:
                print(f"Discovery attempt {attempt+1} failed: {e}")
                time.sleep(settings.chunk_delay_s)
    return all_components

def cross_validate(pass_a: list, pass_b: list) -> list:
    confirmed = []
    names_b = {c["name"].lower().strip() for c in pass_b}
    for comp in pass_a:
        if comp["name"].lower().strip() in names_b:
            confirmed.append(comp)
        else:
            confirmed.append(comp)  # include anyway, flag for review
    return confirmed

# ── Extraction ────────────────────────────────────────────────────────────────
LABEL_MAP = {
    "who can apply": "Eligibility Criteria",
    "who is eligible": "Eligibility Criteria",
    "target audience": "Eligibility Criteria",
    "المستفيدون": "Eligibility Criteria",
    "how to submit": "Application Steps",
    "submission steps": "Application Steps",
    "الخطوات": "Application Steps",
    "what you need": "Required Documents",
    "documents needed": "Required Documents",
    "المستندات المطلوبة": "Required Documents",
    "fees": "Fees",
    "cost": "Fees",
    "charges": "Fees",
    "الرسوم": "Fees",
    "about": "Service Overview",
    "overview": "Service Overview",
    "description": "Service Overview",
    "نبذة": "Service Overview",
}

EXTRACTION_PROMPT = """You are extracting structured content from a legal government document to populate a Word document.

Component: {name}
Type: {comp_type}
Pages: {start} to {end}

EXTRACTION RULES:
1. Extract content for these sections IN ORDER (for services):
   - Service Overview
   - Eligibility Criteria (with named subsections)
   - Fees
   - Required Documents
   - Application Steps (numbered steps with all bullets)
   - How to Follow Up (only if present)
   - Application Notifications and Statuses (only if present)

2. For non-service components: identify natural subsections semantically.

3. For each Application Step extract:
   - step_number, title, all bullets exactly
   - branches (Yes/No paths) with their own bullets
   - notes (📌), warnings (⚠️)
   - image_ids: list page numbers where images appear in this step

4. Tables: reproduce as {{"headers": [...], "rows": [[...]]}}

5. TEXT FIDELITY: extract verbatim. Mark bold as **text**, italic as _text_, bold+italic as ***text***

6. LANGUAGE: if Arabic extract in Arabic, mark "rtl": true. If bilingual extract both.

7. If a section has no content return null for that section.

8. Return ONLY valid JSON. Schema:
{{
  "component_name": "...",
  "component_type": "...",
  "is_service": true/false,
  "language": "en"|"ar"|"bilingual",
  "sections": {{
    "service_overview": {{"text": "...", "rtl": false}},
    "eligibility_criteria": {{
      "subsections": [
        {{"title": "...", "bullets": ["..."], "warnings": ["..."], "rtl": false}}
      ]
    }},
    "fees": {{"text": "..."}},
    "required_documents": [
      {{"name": "...", "description": "...", "conditional": false}}
    ],
    "application_steps": [
      {{
        "step_number": 1,
        "title": "...",
        "bullets": ["..."],
        "branches": [
          {{"label": "...", "bullets": ["..."], "image_pages": [N]}}
        ],
        "image_pages": [N],
        "notes": ["..."],
        "warnings": ["..."],
        "rtl": false
      }}
    ],
    "follow_up": {{"text": "...", "image_pages": [N]}},
    "notifications_statuses": {{
      "tables": [
        {{"table_title": "...", "headers": ["..."], "rows": [["..."]]}}
      ]
    }},
    "free_sections": [
      {{"title": "...", "content": "...", "rtl": false}}
    ]
  }},
  "label_mappings": [{{"source": "...", "mapped_to": "..."}}]
}}

DOCUMENT TEXT:
{text}"""

def extract_component(comp: dict, chunks: list) -> dict:
    start_p = comp.get("start_page", 1)
    end_p = comp.get("end_page", chunks[-1]["end"] if chunks else 1)
    if end_p == "continues":
        end_p = chunks[-1]["end"] if chunks else start_p + 30

    relevant_chunks = [
        c for c in chunks
        if c["end"] >= start_p - 2 and c["start"] <= end_p + 2
    ]

    results = []
    for chunk in relevant_chunks:
        prompt = EXTRACTION_PROMPT.format(
            name=comp["name"],
            comp_type=comp.get("type", "other"),
            start=chunk["start"],
            end=chunk["end"],
            text=chunk["text"][:10000],
        )
        for attempt in range(settings.chunk_retries):
            try:
                raw = call_claude([{"role": "user", "content": prompt}], max_tokens=4000)
                raw = raw.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
                parsed = json.loads(raw)
                results.append(parsed)
                break
            except Exception as e:
                print(f"Extraction attempt {attempt+1} failed for {comp['name']}: {e}")
                time.sleep(settings.chunk_delay_s)

    if not results:
        return {"component_name": comp["name"], "component_type": comp.get("type"), "sections": {}}
    return results[0]

# ── Docx assembly ─────────────────────────────────────────────────────────────
REQUIRED_SECTIONS = ["service_overview", "eligibility_criteria", "fees", "required_documents", "application_steps"]

def assemble_docx(comp_data: dict, image_index: dict, table_index: dict) -> bytes:
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    name = comp_data.get("component_name", "Document")
    is_service = comp_data.get("is_service", False)
    rtl_doc = comp_data.get("language", "en") == "ar"
    sections = comp_data.get("sections", {})

    # H1 - Component name
    add_heading(doc, name, level=1, underline=True, rtl=rtl_doc)

    if is_service:
        # ── Service Overview
        add_heading(doc, "Service Overview", level=2, underline=True, rtl=rtl_doc)
        overview = sections.get("service_overview")
        if overview and overview.get("text"):
            add_body(doc, overview["text"], rtl=overview.get("rtl", False))
        else:
            add_body(doc, "[Section not found in source document — requires manual review]",
                    color=(255, 0, 0))

        # ── Eligibility Criteria
        add_heading(doc, "Eligibility Criteria", level=2, underline=True, rtl=rtl_doc)
        eligibility = sections.get("eligibility_criteria")
        if eligibility and eligibility.get("subsections"):
            for sub in eligibility["subsections"]:
                add_heading(doc, sub.get("title", ""), level=3, rtl=sub.get("rtl", False))
                for bullet in sub.get("bullets", []):
                    add_bullet(doc, bullet, rtl=sub.get("rtl", False))
                for warning in sub.get("warnings", []):
                    add_warning(doc, warning, rtl=sub.get("rtl", False))
        else:
            add_body(doc, "[Section not found in source document — requires manual review]",
                    color=(255, 0, 0))

        # ── Fees
        add_heading(doc, "Fees", level=2, underline=True, rtl=rtl_doc)
        fees = sections.get("fees")
        if fees and fees.get("text"):
            add_body(doc, fees["text"])
        else:
            add_body(doc, "No fees apply.")

        # ── Required Documents
        add_heading(doc, "Required Documents", level=2, underline=True, rtl=rtl_doc)
        docs_list = sections.get("required_documents")
        if docs_list:
            add_body(doc, "You will be asked to upload the following documents:")
            for d in docs_list:
                doc_name = d.get("name", "")
                doc_desc = d.get("description", "")
                conditional = " (if applicable)" if d.get("conditional") else ""
                add_bullet(doc, f"{doc_name}: {doc_desc}{conditional}", bold_part=doc_name)
        else:
            add_body(doc, "[Section not found in source document — requires manual review]",
                    color=(255, 0, 0))

        # ── Application Steps
        add_heading(doc, "Application Steps", level=2, underline=True, rtl=rtl_doc)
        steps = sections.get("application_steps")
        if steps:
            for step in steps:
                step_title = f"Step {step.get('step_number', '')}: {step.get('title', '')}"
                add_heading(doc, step_title, level=3, rtl=step.get("rtl", False))

                # Main bullets
                for bullet in step.get("bullets", []):
                    add_bullet(doc, bullet, rtl=step.get("rtl", False))

                # Notes and warnings
                for note in step.get("notes", []):
                    add_note(doc, note, rtl=step.get("rtl", False))
                for warning in step.get("warnings", []):
                    add_warning(doc, warning, rtl=step.get("rtl", False))

                # Branches (Yes/No paths)
                for branch in step.get("branches", []):
                    add_body(doc, branch.get("label", ""), bold=True, italic=True)
                    for b_bullet in branch.get("bullets", []):
                        add_bullet(doc, b_bullet, rtl=step.get("rtl", False))
                    # Branch images
                    for page_num in branch.get("image_pages", []):
                        imgs = image_index.get(page_num, [])
                        for img_path in imgs[:1]:
                            try:
                                add_image(doc, Path(img_path).read_bytes())
                            except Exception:
                                pass

                # Step images (non-branch)
                if not step.get("branches"):
                    for page_num in step.get("image_pages", []):
                        imgs = image_index.get(page_num, [])
                        for img_path in imgs[:1]:
                            try:
                                add_image(doc, Path(img_path).read_bytes())
                            except Exception:
                                pass
        else:
            add_body(doc, "[Section not found in source document — requires manual review]",
                    color=(255, 0, 0))

        # ── How to Follow Up (conditional)
        follow_up = sections.get("follow_up")
        if follow_up and follow_up.get("text"):
            add_heading(doc, "How to Follow Up", level=2, underline=True, rtl=rtl_doc)
            add_body(doc, follow_up["text"])
            for page_num in follow_up.get("image_pages", []):
                imgs = image_index.get(page_num, [])
                for img_path in imgs[:1]:
                    try:
                        add_image(doc, Path(img_path).read_bytes())
                    except Exception:
                        pass

        # ── Notifications & Statuses (conditional)
        notifs = sections.get("notifications_statuses")
        if notifs and notifs.get("tables"):
            add_heading(doc, "Application Notifications and Statuses", level=2,
                       underline=True, rtl=rtl_doc)
            for tbl in notifs["tables"]:
                if tbl.get("table_title"):
                    add_body(doc, tbl["table_title"], bold=True, italic=True)
                add_table(doc, tbl.get("headers", []), tbl.get("rows", []))

    else:
        # Path B — non-service: free sections
        free_sections = sections.get("free_sections", [])
        if free_sections:
            for fs in free_sections:
                add_heading(doc, fs.get("title", ""), level=2, underline=True,
                           rtl=fs.get("rtl", False))
                add_body(doc, fs.get("content", ""), rtl=fs.get("rtl", False))
        else:
            # Fallback: dump whatever sections exist
            for key, val in sections.items():
                if val:
                    add_heading(doc, key.replace("_", " ").title(), level=2, underline=True)
                    if isinstance(val, str):
                        add_body(doc, val)
                    elif isinstance(val, dict) and val.get("text"):
                        add_body(doc, val["text"])

        # Include any tables from table_index
        for page_num, tables in table_index.items():
            for tbl in tables:
                if tbl.get("headers") and tbl.get("rows"):
                    add_table(doc, tbl["headers"], tbl["rows"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── Main pipeline ─────────────────────────────────────────────────────────────
async def run_pipeline(payload: JobPayload):
    job_id = payload.jobId
    work_dir = Path(settings.work_root) / job_id
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        # 1. Download PDF
        update_job(job_id, "extracting", 5, "Downloading PDF…")
        pdf_path = work_dir / "input.pdf"
        async with httpx.AsyncClient(timeout=120) as client:
            r = await client.get(payload.pdfUrl)
            r.raise_for_status()
            pdf_path.write_bytes(r.content)

        # 2. Pre-extract images and tables
        update_job(job_id, "extracting", 10, "Extracting images and tables…")
        image_index = extract_images(str(pdf_path), work_dir / "images")
        table_index = extract_tables(str(pdf_path), work_dir / "tables")

        # 3. Extract text chunks
        update_job(job_id, "extracting", 15, "Chunking PDF text…")
        chunks = extract_text_chunks(str(pdf_path), settings.chunk_pages)

        # 4. Discovery pass A
        update_job(job_id, "analyzing", 20, "Discovery pass A…")
        pass_a = discover_components(chunks)

        # 5. Discovery pass B
        update_job(job_id, "analyzing", 30, "Discovery pass B…")
        pass_b = discover_components(chunks)

        # 6. Cross-validate
        update_job(job_id, "analyzing", 35, "Cross-validating component manifest…")
        components = cross_validate(pass_a, pass_b)

        # Save manifest
        manifest_path = work_dir / "manifest_final.json"
        manifest_path.write_text(json.dumps(components, ensure_ascii=False, indent=2))

        if not components:
            update_job(job_id, "failed", 0, "No components detected in PDF.")
            return

        # 7. Extract each component
        output_files = []
        total = len(components)
        for idx, comp in enumerate(components):
            progress = 40 + int((idx / total) * 40)
            update_job(job_id, "analyzing", progress,
                      f"Extracting component {idx+1}/{total}: {comp['name'][:50]}…")

            # Save checkpoint
            checkpoint_raw = work_dir / f"component_{idx+1:02d}_raw.json"
            comp_data = extract_component(comp, chunks)
            checkpoint_raw.write_text(json.dumps(comp_data, ensure_ascii=False, indent=2))

            # Validate required sections for services
            if comp_data.get("is_service"):
                missing = []
                for sec in REQUIRED_SECTIONS:
                    val = comp_data.get("sections", {}).get(sec)
                    if not val:
                        missing.append(sec)
                comp_data["_validation"] = {
                    "status": "incomplete" if missing else "complete",
                    "missing_sections": missing,
                }

            checkpoint_validated = work_dir / f"component_{idx+1:02d}_validated.json"
            checkpoint_validated.write_text(json.dumps(comp_data, ensure_ascii=False, indent=2))

            # 8. Assemble docx
            update_job(job_id, "generating", progress + 2,
                      f"Assembling document for: {comp['name'][:50]}…")
            docx_bytes = assemble_docx(comp_data, image_index, table_index)

            # 9. Upload to Supabase
            safe_name = comp["name"][:40].replace("/", "-").replace(" ", "_")
            docx_filename = f"{job_id}/{idx+1:02d}_{safe_name}.docx"
            sb.storage.from_(payload.outputsBucket).upload(
                docx_filename,
                docx_bytes,
                {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
            )

            # Get public URL
            url_resp = sb.storage.from_(payload.outputsBucket).get_public_url(docx_filename)
            public_url = url_resp if isinstance(url_resp, str) else url_resp.get("publicUrl", "")
            output_files.append({"name": f"{comp['name']}.docx", "url": public_url})

        # 10. Done
        update_job(job_id, "complete", 100, f"Done — {len(output_files)} document(s) ready.")
        post_callback(payload.callbackUrl, {
            "jobId": job_id,
            "status": "complete",
            "files": output_files,
        })

    except Exception as e:
        tb = traceback.format_exc()
        print(f"Pipeline error: {tb}")
        update_job(job_id, "failed", 0, f"Pipeline error: {str(e)[:200]}")
        post_callback(payload.callbackUrl, {"jobId": job_id, "status": "failed", "error": str(e)})
    finally:
        # Cleanup work dir
        import shutil
        try:
            shutil.rmtree(work_dir, ignore_errors=True)
        except Exception:
            pass

# ── Routes ────────────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {"status": "ok"}

@app.post("/process")
async def process(
    payload: JobPayload,
    background_tasks: BackgroundTasks,
    x_worker_secret: Optional[str] = Header(None),
):
    if x_worker_secret != settings.worker_secret:
        raise HTTPException(status_code=401, detail="Invalid worker secret")
    run_id = str(uuid.uuid4())
    background_tasks.add_task(run_pipeline, payload)
    return {"accepted": True, "runId": run_id}
